import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document
import bcrypt
from supabase import create_client
import land_development

# --- Supabase Connection ---
url = st.secrets["SUPABASE_URL"]
key = st.secrets["SUPABASE_KEY"]
supabase = create_client(url, key)

# --- Load users from Supabase ---
@st.cache_data
def load_users():
    response = supabase.table("users").select("*").execute()
    data = response.data
    if not data:
        st.warning("⚠️ No users found in Supabase.")
        return pd.DataFrame()
    return pd.DataFrame(data)

# --- Password Check ---
def check_password(username, password, users_df):
    if 'username' not in users_df.columns or 'password_hash' not in users_df.columns:
        st.error("❌ Required columns missing in users table.")
        return False

    user_row = users_df[users_df['username'] == username]
    if user_row.empty:
        return False

    stored_hash = user_row.iloc[0].get("password_hash")
    if not stored_hash:
        return False
    return bcrypt.checkpw(password.encode(), stored_hash.encode())

# --- Password Update ---
def update_password(username, current_password, new_password):
    try:
        users_df = load_users()

        if 'username' not in users_df.columns or 'password_hash' not in users_df.columns:
            st.error("⚠️ Required columns missing in users table.")
            return False

        user_row = users_df[users_df['username'] == username]
        if user_row.empty:
            st.error("❌ User not found.")
            return False

        stored_hash = user_row.iloc[0].get('password_hash')
        if not stored_hash or not bcrypt.checkpw(current_password.encode(), stored_hash.encode()):
            st.error("🔐 Current password is incorrect.")
            return False

        new_hash = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
        supabase.table("users").update({"password_hash": new_hash}).eq("username", username).execute()

        st.success("✅ Password updated successfully.")
        return True

    except Exception as e:
        st.error(f"🚨 An unexpected error occurred: {e}")
        return False

# --- Initialize session state ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.username = ""

# --- Login UI ---
if not st.session_state.logged_in:
    st.title("🔐 Secure Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    users_df = load_users()

    if st.button("Login"):
        if check_password(username, password, users_df):
            st.session_state.logged_in = True
            st.session_state.username = username
            st.success(f"Welcome, {username}!")
            st.rerun()
        else:
            st.error("Invalid username or password.")
    st.stop()


# --- Protected App Content ---
st.set_page_config(page_title="Land Acquisition Tracker", layout="wide")
st.sidebar.image("logo.png", width=150)
st.sidebar.markdown("## Scott Communities")

if st.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.rerun()

# --- Password Update Section ---
with st.sidebar.expander("🔒 Change Password"):
    with st.form("change_password_form"):
        current_pw = st.text_input("Current Password", type="password")
        new_pw = st.text_input("New Password", type="password")
        confirm_pw = st.text_input("Confirm New Password", type="password")
        submit_pw = st.form_submit_button("Update Password")

    if submit_pw:
        if new_pw != confirm_pw:
            st.error("New passwords do not match.")
        elif not new_pw:
            st.error("New password cannot be empty.")
        else:
            update_password(st.session_state.username, current_pw, new_pw)

# --- Constants ---
STATUS_OPTIONS = ["Not started", "Underwriting", "Initial Review", "LOI", "Second Review", "PSA", "No Go"]
PROPERTY_TYPES = ["BTR", "Commercial", "Industrial", "Mixed-use", "Multi-family", "Single-family", "Town Homes"]

# --- Session State ---
if "deal_saved" not in st.session_state:
    st.session_state.deal_saved = False
# --- Sidebar Navigation ---
main_menu = st.sidebar.selectbox("Main Menu", ["Land Acquisition", "Land Development"])
if main_menu == "Land Development":
    land_development.render_land_development_ui()
    st.stop()

# --- Land Acquisition Views ---
st.title("📍 Land Acquisition Tracker")
view = st.sidebar.radio("Select View", ["Search Properties", "Add New Deal", "Preview Deals Table", "LOI Creation"])

# --- Helper Functions ---
def generate_new_pk():
    year = datetime.now().strftime("%Y")
    response = supabase.table("deals").select("pk").execute()
    pks = [r["pk"] for r in response.data if r["pk"].startswith(f"{year}-")]
    if pks:
        last_num = max([int(pk.split("-")[1]) for pk in pks])
        new_num = last_num + 1
    else:
        new_num = 1
    return f"{year}-{new_num:04d}"

def upsert_deal(data):
    pk = data["pk"]
    existing = supabase.table("deals").select("pk").eq("pk", pk).execute()

    for key in ["size", "asking_price", "proposed_price"]:
        if not data[key] or data[key] in ["", "TBD", "N/A"]:
            data[key] = None
    try:
        if existing.data:
            supabase.table("deals").update(data).eq("pk", pk).execute()
        else:
            supabase.table("deals").insert(data).execute()
    except Exception as e:
        st.error(f"❌ Supabase error: {e}")

def get_deal(pk):
    response = supabase.table("deals").select("*").eq("pk", pk).execute()
    return response.data[0] if response.data else None

def is_duplicate_deal(name, location, prop_type, size):
    response = supabase.table("deals").select("pk").match({
        "property_name": name,
        "location": location,
        "property_type": prop_type,
        "size": size
    }).execute()
    return bool(response.data)

def paginate_dataframe(df, items_per_page=5, label="Page"):
    total_items = len(df)
    if total_items == 0:
        st.info("No results found.")
        return df.iloc[0:0], 0
    total_pages = (total_items - 1) // items_per_page + 1
    col1, col2 = st.columns([8, 2])
    with col2:
        page = st.number_input(label, min_value=1, max_value=total_pages, value=1, step=1, label_visibility="collapsed")
    start_idx = (page - 1) * items_per_page
    return df.iloc[start_idx:start_idx + items_per_page], page

def render_status_cell(status):
    color_map = {
        "Not started": ("#B0B0B0", 100),
        "Underwriting": ("#FFF176", 20),
        "Initial Review": ("#FFD54F", 40),
        "LOI": ("#AED581", 60),
        "Second Review": ("#66BB6A", 80),
        "PSA": ("#66BB6A", 100),
        "No Go": ("#FF6961", 100),
    }
    color, fill = color_map.get(status, ("#FFFFFF", 0))
    return f"""
    <div style="
        width: 100%;
        height: 24px;
        border: 1px solid #ccc;
        background: linear-gradient(to right, {color} {fill}%, #E0E0E0 {fill}%);
        text-align: center;
        line-height: 24px;
        font-size: 12px;
        font-weight: 500;
    ">
        <strong>{status}</strong>
    </div>
    """
# --- Search Properties ---
if view == "Search Properties":
    st.header("🔍 Search Properties")
    activity_filter = st.selectbox("Activity", ["Active", "No Go", "All"])
    status_filter = st.selectbox("Status", ["All"] + STATUS_OPTIONS)

    query = supabase.table("deals").select(
        "pk, status, property_name, location, property_type, size, asking_price, proposed_price, link"
    )
    if activity_filter != "All":
        query = query.eq("activity", activity_filter)
    if status_filter != "All":
        query = query.eq("status", status_filter)

    response = query.execute()
    df = pd.DataFrame(response.data)

    if df.empty:
        st.info("No matching records found.")
    else:
        df["_sort"] = df["status"].map(lambda x: STATUS_OPTIONS.index(x) if x in STATUS_OPTIONS else 99)
        df = df.sort_values(by=["_sort", "property_name"]).drop(columns=["_sort"])
        df_page, page = paginate_dataframe(df)

        headers = ["Status", "Property Name", "Location", "Property Type", "Size", "Asking Price", "Proposed Price", "Link", "Edit"]
        for col, header in zip(st.columns([1.2, 1.5, 1.5, 1.2, 1.2, 1.2, 1.2, 0.8, 0.8]), headers):
            col.markdown(f"**{header}**")

        for _, row in df_page.iterrows():
            deal = get_deal(row["pk"])
            link_url = deal.get("link", "")
            c = st.columns([1.2, 1.5, 1.5, 1.2, 1.2, 1.2, 1.2, 0.8, 0.8])
            c[0].markdown(render_status_cell(row["status"]), unsafe_allow_html=True)
            c[1].markdown(f"<strong>{row['property_name']}</strong>", unsafe_allow_html=True)
            c[2].write(row["location"])
            c[3].write(row["property_type"])
            c[4].write(row["size"])
            c[5].write(row["asking_price"] or "TBD")
            c[6].write(row["proposed_price"] or "TBD")
            c[7].markdown(f'<a href="{link_url}" target="_blank">link</a>' if link_url else "TBD", unsafe_allow_html=True)
            if c[8].button("✏️ Edit", key=f"edit-{row['pk']}"):
                st.session_state.edit_pk = row["pk"]

        if "edit_pk" in st.session_state:
            deal = get_deal(st.session_state.edit_pk)
            if deal:
                st.subheader(f"📝 Edit Deal: {deal['pk']}")
                with st.form("edit_form"):
                    status = st.selectbox("Status", STATUS_OPTIONS, index=STATUS_OPTIONS.index(deal["status"]))
                    name = st.text_input("Property Name", value=deal["property_name"])
                    location = st.text_input("Location", value=deal["location"])
                    prop_type = st.text_input("Property Type", value=deal["property_type"])
                    size = st.text_input("Size", value=deal["size"])
                    asking = st.text_input("Asking Price", value=deal["asking_price"])
                    proposed = st.text_input("Proposed Price", value=deal["proposed_price"])
                    link = st.text_input("Link", value=deal["link"])
                    receive_dt = deal["receive_dt"]
                    st.markdown(f"**Receive Date:** `{receive_dt}`")
                    activity = "No Go" if status == "No Go" else "Active"

                    colA, colB = st.columns([1, 1])
                    save = colA.form_submit_button("💾 Save Changes")
                    cancel = colB.form_submit_button("❌ Cancel")

                if save:
                    now = datetime.now().isoformat()
                    updated = {
                        "pk": deal["pk"],
                        "status": status,
                        "property_name": name,
                        "location": location,
                        "property_type": prop_type,
                        "size": size,
                        "asking_price": asking,
                        "proposed_price": proposed,
                        "link": link,
                        "receive_dt": receive_dt,
                        "underwriting_dt": deal["underwriting_dt"] or (now if status == "Underwriting" and deal["status"] != "Underwriting" else None),
                        "initial_review_dt": deal["initial_review_dt"] or (now if status == "Initial Review" and deal["status"] != "Initial Review" else None),
                        "loi_dt": deal["loi_dt"] or (now if status == "LOI" and deal["status"] != "LOI" else None),
                        "second_review_dt": deal["second_review_dt"] or (now if status == "Second Review" and deal["status"] != "Second Review" else None),
                        "psa_dt": deal["psa_dt"] or (now if status == "PSA" and deal["status"] != "PSA" else None),
                        "no_go_dt": deal["no_go_dt"] or (now if status == "No Go" and deal["status"] != "No Go" else None),
                        "activity": activity
                    }
                    upsert_deal(updated)
                    st.success("✅ Changes saved.")
                    del st.session_state.edit_pk
                    st.rerun()
                elif cancel:
                    del st.session_state.edit_pk
                    st.info("🛑 Edit canceled.")
                    st.rerun()
# --- Add New Deal ---
elif view == "Add New Deal":
    st.header("➕ Add New Deal")
    if st.session_state.deal_saved:
        st.success("🎉 Deal added successfully.")
        st.session_state.deal_saved = False
        st.stop()

    with st.form("add_deal_form"):
        pk = generate_new_pk()
        st.markdown(f"**New Deal ID:** `{pk}`")
        name = st.text_input("Property Name")
        location = st.text_input("Location")
        prop_type = st.selectbox("Property Type", PROPERTY_TYPES)
        size = st.text_input("Size")
        asking = st.text_input("Asking Price")
        proposed = st.text_input("Proposed Price")
        link = st.text_input("Link to File")
        receive_dt = datetime.now().isoformat()
        st.markdown(f"**Receive Date:** `{receive_dt}`")
        status = "Not started"
        activity = "Active"

        submit = st.form_submit_button("✅ Save Deal")
        if submit:
            if not name or not location:
                st.error("❗ Property Name and Location are required.")
            elif is_duplicate_deal(name, location, prop_type, size):
                st.warning("⚠️ A deal with the same Property Name, Location, Property Type, and Size already exists.")
            else:
                data = {
                    "pk": pk,
                    "status": status,
                    "property_name": name,
                    "location": location,
                    "property_type": prop_type,
                    "size": size,
                    "asking_price": asking,
                    "proposed_price": proposed,
                    "link": link,
                    "receive_dt": receive_dt,
                    "underwriting_dt": None,
                    "initial_review_dt": None,
                    "loi_dt": None,
                    "second_review_dt": None,
                    "psa_dt": None,
                    "no_go_dt": None,
                    "activity": activity
                }
                upsert_deal(data)
                st.session_state.deal_saved = True
                st.rerun()

# --- Preview Deals Table ---
elif view == "Preview Deals Table":
    st.header("📊 Full Deal Inventory")
    response = supabase.table("deals").select("*").execute()
    full_df = pd.DataFrame(response.data)
    st.dataframe(full_df)

# --- LOI Creation ---
elif view == "LOI Creation":
    st.header("📄 LOI Creation")
    response = supabase.table("deals").select("pk, property_name").in_("status", ["LOI", "Second Review", "PSA"]).order("pk", desc=True).execute()
    deals = response.data
    deal_options = {f"{d['pk']} - {d['property_name']}": d["pk"] for d in deals}

    if not deal_options:
        st.info("No eligible deals found for LOI creation.")
    else:
        with st.form("loi_form"):
            st.subheader("Fill in LOI Details")
            selected_deal = st.selectbox("Select Deal", list(deal_options.keys()))
            deal_pk = deal_options.get(selected_deal)

            date_of_letter = st.date_input("Date of Letter", value=datetime.today()).strftime("%B %d, %Y")
            seller_name = st.text_input("Seller Name")
            company_name_of_seller = st.text_input("Company Name of Seller")
            broker_name = st.text_input("Broker Name")
            broker_address1 = st.text_input("Broker Address Line 1")
            broker_address2 = st.text_input("Broker Address Line 2")
            property_full_description = st.text_area("Property Name with Address and Description")
            purchase_price = st.text_input("Purchase Price")
            earnest_money_1 = st.text_input("Earnest Money (Initial Deposit)")
            earnest_money_2 = st.text_input("Earnest Money (Second Deposit)")
            feasibility_period = st.text_input("Feasibility Period")
            close_of_escrow = st.text_input("Close of Escrow Period")
            signer_name = st.text_input("Signer Name")
            signer_title = st.text_input("Signer Title")
            sent_at = st.date_input("Date Sent (optional)", value=None)
            received_at = st.date_input("Date Received (optional)", value=None)
            notes = st.text_area("Notes (optional)")

            submit = st.form_submit_button("📝 Generate LOI")

        if submit and deal_pk:
            def replace_placeholder_in_paragraph(paragraph, replacements):
                full_text = ''.join(run.text for run in paragraph.runs)
                for old, new in replacements.items():
                    full_text = full_text.replace(old, new)
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                else:
                    paragraph.add_run(full_text)

            def replace_placeholders(doc, replacements):
                for paragraph in doc.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, replacements)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_placeholder_in_paragraph(paragraph, replacements)
                for section in doc.sections:
                    for paragraph in section.header.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, replacements)
                    for paragraph in section.footer.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, replacements)

            template_path = "LOI_Template_Do_Not_Remove.docx"
            try:
                doc = Document(template_path)
            except Exception as e:
                st.error(f"❌ Failed to load template: {e}")
                st.stop()

            replacements = {
                "{{DATE_OF_LETTER}}": date_of_letter,
                "{{SELLER_NAME}}": seller_name,
                "{{COMPANY_NAME_OF_SELLER}}": company_name_of_seller,
                "{{BROKER_NAME}}": broker_name,
                "{{BROKER_ADDRESS_LINE_1}}": broker_address1,
                "{{BROKER_ADDRESS_LINE_2}}": broker_address2,
                "{{SELLER_NAME_WITH_PREFIX}}": f"Mr./Ms. {seller_name}",
                "{{PROPERTY NAME_WITH_ADDRESS_AND_DESCRIPTION_(See _attachment)}}": property_full_description,
                "{{PURCHASE_PRICE}}": purchase_price,
                "{{EARNEST_MONEY_1}}": earnest_money_1,
                "{{EARNEST_MONEY_2}}": earnest_money_2,
                "{{FEASIBILITY_PERIOD}}": feasibility_period,
                "{{CLOSE_OF_ESCROW}}": close_of_escrow,
                "{{SIGNER_NAME}}": signer_name,
                "{{SIGNER_TITLE}}": signer_title
            }

            replace_placeholders(doc, replacements)

            deal_info = supabase.table("deals").select("property_name, property_type").eq("pk", deal_pk).execute().data[0]
            property_name = deal_info["property_name"]
            property_type = deal_info["property_type"]

            def sanitize(text):
                return text.strip().replace(" ", "_").replace("/", "-")

            safe_name = sanitize(property_name)
            safe_type = sanitize(property_type)
            filename = f"LOI_{safe_name}_{safe_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            supabase.table("lois").insert({
                "deal_pk": deal_pk,
                "created_at": datetime.now().isoformat(),
                "sent_at": sent_at.strftime("%Y-%m-%d") if sent_at else None,
                "received_at": received_at.strftime("%Y-%m-%d") if received_at else None,
                "file_path": filename,
                "notes": notes
            }).execute()

            st.success("✅ LOI generated.")
            st.download_button("📥 Download LOI", doc_io, file_name=filename)
